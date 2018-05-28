# -*- coding: UTF-8 -*-
import  csv
import  sys
import  importlib
import docx
import  os
importlib.reload(sys)
from docx import Document
from pdfminer.pdfparser import PDFParser ,PDFDocument
from pdfminer.pdfinterp import PDFResourceManager,PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout    import LTTextBoxHorizontal,LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed


class DealFile(object):

    def readCsv(self,path):
        infolist = []
        with open(path, "r") as f:
            allFileInfo = csv.reader(f)
            for row in allFileInfo:
                infolist.append(row)
                # print(row)
            # print(allFileInfo)

    def writeCsv(self,path, data):
        with open(path, "w")as f:
            writer = csv.writer(f)
            for rowData in data:
                writer.writerow(rowData)

    def readPDF(self,path,callBlcak=None,toPath=""):
        f = open(path, "rb")

        # 创建一个PDF分析器
        parse = PDFParser(f)

        # 创建PDF文件
        pdfFile = PDFDocument()

        # 链接
        parse.set_document(pdfFile)
        # 双向链接
        pdfFile.set_parser(parse)

        # 提供初始化密码
        pdfFile.initialize()

        # 检测文档是否提供TXT转换
        if not pdfFile.is_extractable:
            # 没有结束了
            raise PDFTextExtractionNotAllowed
        else:
            # 解析数据
            # 数据管理器
            manager = PDFResourceManager()
            # PDF设备管理对象
            laparam = LAParams()
            device = PDFPageAggregator(manager, laparams=laparam)
            # 解释器对象
            interpreter = PDFPageInterpreter(manager, device)
            # 开始循环处理    每次处理一页
            for page in pdfFile.get_pages():
                interpreter.process_page(page)
                layout = device.get_result()
                # 处理图册
                for x in layout:
                    if (isinstance(x, LTTextBoxHorizontal)):
                        str = x.get_text()
                        if toPath=="":
                            if callBlcak!=None:
                                 callBlcak(str)
                            else:
                                 print(str)
                        else:
                            with open(toPath, "a")as f:
                                print(str)
                                # 处理数据、
                                f.write(str + "\n")

    def readWord(self,path):
        document = Document(path)
        for paragraph in document.paragraphs:
            print(paragraph.text)

if __name__ == '__main__':
    d = DealFile()
    path ="/Users/zzs/Desktop/python/校趣标注.pdf"
    def func(str):
        print(str+"!!!!!!!")
    d.readPDF(path,func)
