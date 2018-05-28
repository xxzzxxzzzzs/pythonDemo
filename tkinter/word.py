# -*- coding: cp936 -*-
# -*- coding: UTF-8 -*-
import docx
import  os

from docx import Document
path="/Users/zzs/Documents/关于Git客户端的用法.docx"
def readWord(path):
    document = Document(path)
    for paragraph in document.paragraphs:
        print(paragraph.text)
def makeWord(names):
    for name in  names:
        path=os.path.join(os.getcwd(),name)
        doc=docx.Document()
        doc.add_paragraph("啊")
        doc.add_paragraph("山，真的很高"+name+"    ")
        doc.add_paragraph("高的吓人")
        doc.add_paragraph("啊！！")
        doc.add_paragraph("实在是高")
        doc.save(path+".docx")




# readWord(path)

makeWord(["zzs1","zzs2","zzs3"])


